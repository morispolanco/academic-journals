import streamlit as st
from docx import Document
import re
from io import StringIO
import spacy

# Load spaCy model for keyword extraction
nlp = spacy.load("en_core_web_sm")

# Configuración inicial de la página
st.set_page_config(page_title="Manuscript Review Assistant", page_icon="📚")

# Título y descripción
st.title("Manuscript Review Assistant")
st.write("""
Welcome to the Manuscript Review Assistant! This tool helps authors determine if their manuscript meets the guidelines of top management journals. 
Upload your manuscript, select the target journal, and receive personalized feedback.
""")

# Menú en la barra lateral para seleccionar la revista
st.sidebar.header("Select a Journal")
journal = st.sidebar.selectbox(
    "Choose the journal you are targeting:",
    [
        "Academy of Management Journal (AMJ)",
        "Administrative Science Quarterly (ASQ)",
        "Strategic Management Journal (SMJ)",
        "Journal of Management (JOM)",
        "Organization Science"
    ]
)

# Cargar archivo Word
uploaded_file = st.file_uploader("Upload your manuscript (Word file)", type=["docx"])
if uploaded_file:
    try:
        document = Document(uploaded_file)
        st.success("File uploaded successfully!")
        
        # Extraer texto del manuscrito con secciones
        def extract_sections(doc):
            sections = {}
            current_section = None
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    current_section = para.text.strip()
                    sections[current_section] = []
                elif current_section:
                    sections[current_section].append(para.text)
            return {k: "\n".join(v) for k, v in sections.items()}
        
        sections = extract_sections(document)
        manuscript_text = "\n".join([para.text for para in document.paragraphs])
        
        st.subheader("Manuscript Preview")
        st.text_area("Manuscript Content", manuscript_text, height=300)
        
        # Verificar formato básico
        def check_format(text):
            word_count = len(re.findall(r'\w+', text))
            pages_estimate = word_count / 250  # Estimación: 250 palabras por página
            return {
                "word_count": word_count,
                "pages_estimate": pages_estimate
            }
        
        format_info = check_format(manuscript_text)
        
        # Feedback general
        st.subheader("General Feedback")
        st.write(f"- Estimated word count: **{format_info['word_count']}**")
        st.write(f"- Estimated page count: **{format_info['pages_estimate']:.1f} pages**")
        
        # Evaluar según la revista seleccionada
        if st.button("Evaluate Manuscript"):
            st.subheader(f"Evaluation for {journal}")
            
            # Abstract validation
            abstract_text = sections.get("Abstract", "")
            abstract_word_count = len(re.findall(r'\w+', abstract_text))
            st.write(f"- Abstract word count: **{abstract_word_count}**")
            
            if journal == "Academy of Management Journal (AMJ)":
                st.write("""
- **Focus and Scope**: Your manuscript should address topics like organizational behavior, leadership, strategy, HR, innovation, or entrepreneurship.
- **Abstract**: Should be ≤ 150 words. Check if your abstract is concise and highlights the research problem, methods, findings, and contributions.
- **Length**: Maximum 40 pages. Your manuscript is estimated to be **{format_info['pages_estimate']:.1f} pages**.
                """.format(format_info=format_info))
                if abstract_word_count > 150:
                    st.warning("Your abstract exceeds the recommended length of 150 words.")
                if format_info['pages_estimate'] > 40:
                    st.warning("Your manuscript exceeds the recommended length. Consider reducing content.")
            
            # Similar evaluations for other journals...
            
            # Keyword extraction
            def extract_keywords(text):
                doc = nlp(text)
                keywords = [chunk.text for chunk in doc.noun_chunks]
                return list(set(keywords))[:5]  # Top 5 unique keywords
            
            keywords = extract_keywords(manuscript_text)
            st.write(f"- Suggested Keywords: **{', '.join(keywords)}**")
            
            # Additional Suggestions
            st.subheader("Additional Suggestions")
            st.write("""
- **Abstract**: Ensure it is clear, concise, and highlights the key contributions of your research.
- **Keywords**: Include 3–5 relevant keywords that reflect the core themes of your manuscript.
- **Formatting**: Use Times New Roman, size 12, double-spaced, with 1-inch margins.
- **References**: Follow APA (7th edition) style for citations and references.
            """)
        
        # Save Evaluation Results
        def generate_report(feedback):
            buffer = StringIO()
            buffer.write("Manuscript Evaluation Report\n")
            buffer.write("=" * 30 + "\n")
            buffer.write(feedback)
            buffer.seek(0)
            return buffer.getvalue()
        
        if st.button("Download Report"):
            feedback = f"Evaluation for {journal}\n"
            feedback += f"- Estimated word count: {format_info['word_count']}\n"
            feedback += f"- Estimated page count: {format_info['pages_estimate']:.1f} pages\n"
            feedback += f"- Suggested Keywords: {', '.join(keywords)}\n"
            st.download_button("Download Report", generate_report(feedback), file_name="evaluation_report.txt")
    
    except Exception as e:
        st.error(f"Error processing file: {e}")
        st.stop()

# Recursos adicionales
st.sidebar.header("Additional Resources")
st.sidebar.write("""
For more details, visit the official websites of the respective journals:
- [AMJ](https://journals.aom.org/amj)
- [ASQ](https://journals.sagepub.com/home/asq)
- [SMJ](https://onlinelibrary.wiley.com/journal/10970266)
- [JOM](https://journals.sagepub.com/home/jom)
- [Organization Science](https://pubsonline.informs.org/journal/orgsci)
""")
